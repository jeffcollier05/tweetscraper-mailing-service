import snscrape.modules.twitter as sntwitter
import os
import docx
from PIL import Image
import time
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import hyperlink
import settings as stt
import emailgenerator as egen
from datetime import date
import re


# Returns list of files in folder
def update_filelist():
    return os.listdir()


class tweetData:
    def __init__(self):
        self.timestamps = []
        self.urls = []
        self.tweetCount = 0

    def gather_tweets(self):
        try:
            # Creates query from settings
            query = '(from:%s) since:%s' % (stt.twitterHandle, stt.timeBoundary)
            
            # Gets tweet query
            for tweet in sntwitter.TwitterSearchScraper(query).get_items():
                if self.tweetCount == int(stt.tweetLimit):
                    break
                else:
                    # Captures tweet picture
                    command = "tweetcapture -sp -n 1 %s" % tweet.url
                    os.system(command)
                    
                    # Generating timestamp of tweet
                    date = str(tweet.date)
                    year = date[:4]
                    month = date[5:7]
                    day = date[8:10]
                    timeHr = date[11:13]
                    timeMin = date[14:16]
                    tweetTimestamp = '%s/%s/%s AT %s:%s UTC' % (month, day, year, timeHr, timeMin)
                    
                    # Adds timestamp and url to list
                    self.timestamps.insert(0, tweetTimestamp)
                    self.urls.insert(0, tweet.url)

                    # Counts one more tweet
                    self.tweetCount +=1
                    print('Tweet downloaded.')
            print('Captured a total of %s tweets from %s.' % (self.tweetCount, stt.twitterHandle))
        except Exception as e:
            raise Exception(f"Error gathering tweets: {e}")


def create_document(tweetData):
    try:
        # Creates instance of document
        doc = docx.Document()
        filelist = update_filelist()
        
        # Document initial formatting
        paragraph_format = doc.styles['Normal'].paragraph_format
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)
        docWidth = 8.5 #inches

        # Creates first page serving as header
        header = doc.add_heading('%s\'s Newsletter' % stt.twitterHandle, 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.sections[0].page_height = Inches(1.5)
        
        # Creates body of newsletter
        i = 0
        for file in filelist[:]:
            if file.endswith(".png"):
                # Adds page
                doc.add_section()
                
                # Sets individual page height
                image = Image.open(file)
                scaleFactor = docWidth / (image.width / 96)
                pageHeight = (image.height * scaleFactor) / 96 + 1
                section = doc.sections[i+1]
                section.page_height = Inches(pageHeight)
                
                # Populates url, timestamp, and tweet
                p = doc.add_paragraph('')
                hyperlink.add_hyperlink(p, tweetData.urls[i], tweetData.urls[i], 'blue', True)
                doc.add_paragraph(tweetData.timestamps[i])
                doc.add_picture(file, width=Inches(docWidth))
                i+=1

        # Formats margins for all pages
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(0.0)
            section.right_margin = Inches(0.0)
            section.top_margin = Inches(0.2)
            section.bottom_margin = Inches(0.0)

        # Saves docx in folder
        doc.save('%s_%s.docx' % (stt.twitterHandle, date.today()))
        print('Saved the newsletter document.')
    except Exception as e:
        raise Exception(f"Error generating newsletter document: {e}")


def clean_folder():
    try:
        # Deletes the png photos and docx in folder
        filelist = update_filelist()
        for file in filelist[:]:
            if ((file.startswith(f"@{stt.twitterHandle}_") and file.endswith("_tweetcapture.png")) 
            or file == (f"{stt.twitterHandle}_{date.today()}.docx")):
                os.remove(file)
                print('Cleaned up folder by removing %s.' % file)
    except Exception as e:
        print(f"Error cleaning up photos from folder: {e}")


if __name__ == "__main__":
    try:
        data = tweetData()
        data.gather_tweets()
        time.sleep(2)
        create_document(data)
        time.sleep(2)
        egen.send_mail()
    except Exception as e:
        print(e)
        print("... aborting script.")
    finally:
        clean_folder()
        print("Program is finished.\n")